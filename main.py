"""
OCR App — FastAPI Backend (single file)
Run with: uvicorn main:app --reload --host 0.0.0.0 --port 8000
"""

from __future__ import annotations

import asyncio
import io
import json
import logging
import mimetypes
import os
import re
import shutil
import tempfile
import time
import uuid
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Any, AsyncIterator, Optional

import aiofiles
from dotenv import load_dotenv
from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    UploadFile,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from loguru import logger
from PIL import Image
from pydantic import BaseModel, Field
from sse_starlette.sse import EventSourceResponse

load_dotenv()

# ─────────────────────────────────────────────────────────────────────────────
# Config
# ─────────────────────────────────────────────────────────────────────────────

ENABLE_DB: bool = os.getenv("ENABLE_DB", "false").lower() in ("true", "1", "yes")
DATABASE_URL: str = os.getenv("DATABASE_URL", "sqlite+aiosqlite:///./ocr_history.db")
MODELS_DIR: Path = Path(os.getenv("MODELS_DIR", "./models_cache")).resolve()
UPLOAD_DIR: Path = Path(os.getenv("UPLOAD_DIR", "./uploads")).resolve()
OUTPUT_DIR: Path = Path(os.getenv("OUTPUT_DIR", "./outputs")).resolve()
HF_TOKEN: Optional[str] = os.getenv("HF_TOKEN") or None
MAX_UPLOAD_MB: int = int(os.getenv("MAX_UPLOAD_MB", "200"))
MODELS_JSON: Path = Path(__file__).parent / "models.json"

for d in (MODELS_DIR, UPLOAD_DIR, OUTPUT_DIR):
    d.mkdir(parents=True, exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# Database (optional)
# ─────────────────────────────────────────────────────────────────────────────

if ENABLE_DB:
    from sqlalchemy import Column, DateTime, Integer, String, Text
    from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
    from sqlalchemy.orm import DeclarativeBase, sessionmaker

    # Ensure aiosqlite prefix for sqlite URLs
    _db_url = DATABASE_URL
    if _db_url.startswith("sqlite:///") and "aiosqlite" not in _db_url:
        _db_url = _db_url.replace("sqlite:///", "sqlite+aiosqlite:///", 1)

    engine = create_async_engine(_db_url, echo=False)
    AsyncSessionLocal = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

    class Base(DeclarativeBase):
        pass

    class JobRecord(Base):
        __tablename__ = "jobs"
        id = Column(Integer, primary_key=True, index=True)
        job_id = Column(String(64), unique=True, index=True)
        status = Column(String(32), default="pending")
        model_id = Column(String(128))
        input_files = Column(Text)   # JSON list
        output_formats = Column(Text)  # JSON list
        created_at = Column(DateTime, default=datetime.utcnow)
        completed_at = Column(DateTime, nullable=True)
        result_paths = Column(Text, nullable=True)  # JSON list
        error = Column(Text, nullable=True)

    async def get_db() -> AsyncIterator[AsyncSession]:
        async with AsyncSessionLocal() as session:
            yield session

    async def init_db() -> None:
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)

else:
    async def init_db() -> None:  # type: ignore[misc]
        pass

    async def get_db():  # type: ignore[misc]
        yield None

# ─────────────────────────────────────────────────────────────────────────────
# In-memory job store
# ─────────────────────────────────────────────────────────────────────────────

_jobs: dict[str, dict[str, Any]] = {}
_download_progress: dict[str, dict[str, Any]] = {}


def _new_job(job_id: str, model_id: str, output_formats: list[str]) -> dict:
    job = {
        "job_id": job_id,
        "status": "pending",
        "model_id": model_id,
        "output_formats": output_formats,
        "progress": 0,
        "message": "Queued",
        "result_paths": [],
        "error": None,
        "created_at": datetime.utcnow().isoformat(),
        "pages": [],
    }
    _jobs[job_id] = job
    return job


# ─────────────────────────────────────────────────────────────────────────────
# Model registry helpers
# ─────────────────────────────────────────────────────────────────────────────

def load_models() -> list[dict]:
    with open(MODELS_JSON) as f:
        data = json.load(f)
    models = data["models"]
    # Patch downloaded status from disk
    for m in models:
        model_path = MODELS_DIR / m["id"]
        m["downloaded"] = model_path.exists() and any(model_path.iterdir()) if model_path.exists() else False
    return models


def save_models(models: list[dict]) -> None:
    with open(MODELS_JSON, "w") as f:
        json.dump({"models": models}, f, indent=2)


def get_model(model_id: str) -> Optional[dict]:
    for m in load_models():
        if m["id"] == model_id:
            return m
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Model download
# ─────────────────────────────────────────────────────────────────────────────

async def download_model(model_id: str) -> None:
    """Download model from HuggingFace with progress updates."""
    model = get_model(model_id)
    if not model:
        raise ValueError(f"Unknown model: {model_id}")

    _download_progress[model_id] = {
        "status": "downloading",
        "progress": 0,
        "message": "Starting download…",
        "model_id": model_id,
    }

    dest = MODELS_DIR / model_id
    dest.mkdir(parents=True, exist_ok=True)

    hf_tag = model["hf_tag"]
    token = HF_TOKEN

    try:
        from huggingface_hub import snapshot_download
        from huggingface_hub.utils import HfHubHTTPError

        _download_progress[model_id]["message"] = f"Downloading {hf_tag} …"
        _download_progress[model_id]["progress"] = 5

        def _sync_download():
            kwargs: dict[str, Any] = {
                "repo_id": hf_tag,
                "local_dir": str(dest),
                "local_dir_use_symlinks": False,
            }
            if token:
                kwargs["token"] = token
            return snapshot_download(**kwargs)

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _sync_download)

        _download_progress[model_id] = {
            "status": "done",
            "progress": 100,
            "message": "Download complete",
            "model_id": model_id,
        }
        logger.info(f"Model {model_id} downloaded to {dest}")

    except Exception as exc:
        logger.error(f"Download failed for {model_id}: {exc}")
        _download_progress[model_id] = {
            "status": "error",
            "progress": 0,
            "message": str(exc),
            "model_id": model_id,
        }
        raise


# ─────────────────────────────────────────────────────────────────────────────
# Image extraction helpers
# ─────────────────────────────────────────────────────────────────────────────

def _extract_images_from_pdf(pdf_path: Path) -> list[Image.Image]:
    from pdf2image import convert_from_path
    return convert_from_path(str(pdf_path), dpi=200)


def _extract_images_from_zip(zip_path: Path) -> list[tuple[str, Image.Image]]:
    results = []
    with zipfile.ZipFile(zip_path) as zf:
        for name in sorted(zf.namelist()):
            lower = name.lower()
            if any(lower.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif")):
                with zf.open(name) as f:
                    img = Image.open(io.BytesIO(f.read())).convert("RGB")
                    results.append((name, img))
    return results


def _load_image(path: Path) -> Image.Image:
    return Image.open(path).convert("RGB")


# ─────────────────────────────────────────────────────────────────────────────
# Element classification helpers
# ─────────────────────────────────────────────────────────────────────────────

def _classify_elements(text: str) -> str:
    """
    Post-process raw model output to tag element types.
    Models like GOT-OCR2 and Florence return structured text;
    we normalise into a common internal format.
    """
    # Already has markdown tables → keep
    # Detect table-like lines
    lines = text.split("\n")
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # Simple heuristic: pipe-separated = table row
        if re.match(r"\|.+\|", line.strip()):
            out.append(line)
        else:
            out.append(line)
        i += 1
    return "\n".join(out)


# ─────────────────────────────────────────────────────────────────────────────
# Model inference
# ─────────────────────────────────────────────────────────────────────────────

_loaded_models: dict[str, Any] = {}


def _get_or_load_model(model_id: str) -> Any:
    """Lazy-load model into memory. Returns (processor, model) tuple."""
    if model_id in _loaded_models:
        return _loaded_models[model_id]

    model_info = get_model(model_id)
    if not model_info:
        raise ValueError(f"Unknown model: {model_id}")

    model_dir = MODELS_DIR / model_id
    if not model_dir.exists():
        raise RuntimeError(f"Model {model_id} not downloaded. Download it first.")

    hf_tag = str(model_dir)  # Use local path
    mtype = model_info.get("type", "vlm")

    import torch
    device = "cuda" if torch.cuda.is_available() else "cpu"
    logger.info(f"Loading model {model_id} on {device}")

    if model_id == "surya-ocr":
        # Surya has its own API
        from surya.model.detection.model import load_model as load_det
        from surya.model.detection.processor import load_processor as load_det_proc
        det_processor, det_model = load_det_proc(), load_det()
        _loaded_models[model_id] = ("surya", det_processor, det_model)
        return _loaded_models[model_id]

    if model_id == "nougat":
        from transformers import NougatProcessor, VisionEncoderDecoderModel
        processor = NougatProcessor.from_pretrained(hf_tag)
        model = VisionEncoderDecoderModel.from_pretrained(hf_tag).to(device)
        _loaded_models[model_id] = ("nougat", processor, model, device)
        return _loaded_models[model_id]

    if model_id in ("got-ocr2",):
        from transformers import AutoModel, AutoTokenizer
        tokenizer = AutoTokenizer.from_pretrained(hf_tag, trust_remote_code=True)
        model = AutoModel.from_pretrained(
            hf_tag, trust_remote_code=True,
            low_cpu_mem_usage=True, device_map="auto" if device == "cuda" else None,
            use_safetensors=True, pad_token_id=tokenizer.eos_token_id
        )
        if device == "cpu":
            model = model.to(device)
        model.eval()
        _loaded_models[model_id] = ("got", tokenizer, model, device)
        return _loaded_models[model_id]

    if model_id == "florence-2":
        from transformers import AutoModelForCausalLM, AutoProcessor
        processor = AutoProcessor.from_pretrained(hf_tag, trust_remote_code=True)
        model = AutoModelForCausalLM.from_pretrained(
            hf_tag, trust_remote_code=True,
            torch_dtype="auto", device_map="auto" if device == "cuda" else None,
        )
        if device == "cpu":
            model = model.to(device)
        _loaded_models[model_id] = ("florence", processor, model, device)
        return _loaded_models[model_id]

    # Default: generic VLM (DeepSeek VL, Qwen2-VL, PaddleOCR-VL, etc.)
    from transformers import AutoModelForCausalLM, AutoProcessor, AutoTokenizer
    try:
        processor = AutoProcessor.from_pretrained(hf_tag, trust_remote_code=True)
    except Exception:
        processor = AutoTokenizer.from_pretrained(hf_tag, trust_remote_code=True)

    model = AutoModelForCausalLM.from_pretrained(
        hf_tag,
        trust_remote_code=True,
        torch_dtype="auto",
        device_map="auto" if device == "cuda" else None,
        low_cpu_mem_usage=True,
    )
    if device == "cpu":
        model = model.to(device)
    model.eval()
    _loaded_models[model_id] = ("generic_vlm", processor, model, device)
    return _loaded_models[model_id]


OCR_PROMPT = (
    "You are an expert OCR engine. Analyze this image carefully and extract ALL content. "
    "Rules:\n"
    "- Preserve ALL text exactly as written\n"
    "- Represent tables as Markdown tables (| col | col |)\n"
    "- Mark diagrams/charts as: [DIAGRAM: brief description]\n"
    "- Mark embedded images/photos as: [IMAGE: brief description]\n"
    "- Mark equations/formulas in LaTeX: $formula$\n"
    "- Preserve heading hierarchy with # ## ###\n"
    "- Preserve bullet lists and numbering\n"
    "- Preserve bold/italic where obvious\n"
    "Output only the extracted content, no commentary."
)


def _run_inference(model_id: str, image: Image.Image) -> str:
    """Run OCR inference on a single PIL image. Returns markdown string."""
    import torch

    loaded = _get_or_load_model(model_id)
    kind = loaded[0]

    if kind == "surya":
        _, processor, model = loaded
        from surya.ocr import run_ocr
        from surya.model.recognition.model import load_model as load_rec
        from surya.model.recognition.processor import load_processor as load_rec_proc
        rec_model = load_rec()
        rec_processor = load_rec_proc()
        langs = [["en"]]
        result = run_ocr([image], langs, model, processor, rec_model, rec_processor)
        lines = [line.text for page in result for line in page.text_lines]
        return "\n".join(lines)

    if kind == "nougat":
        _, processor, model, device = loaded
        pixel_values = processor(image, return_tensors="pt").pixel_values.to(device)
        outputs = model.generate(
            pixel_values,
            min_length=1,
            max_new_tokens=3000,
            bad_words_ids=[[processor.tokenizer.unk_token_id]],
        )
        return processor.batch_decode(outputs, skip_special_tokens=True)[0]

    if kind == "got":
        _, tokenizer, model, device = loaded
        # GOT-OCR2 uses its own chat interface
        result = model.chat(tokenizer, str(image.filename) if hasattr(image, "filename") else "", ocr_type="format")
        return result

    if kind == "florence":
        _, processor, model, device = loaded
        inputs = processor(text="<OCR_WITH_REGION>", images=image, return_tensors="pt").to(device)
        generated_ids = model.generate(
            input_ids=inputs["input_ids"],
            pixel_values=inputs["pixel_values"],
            max_new_tokens=4096,
            num_beams=3,
        )
        result = processor.batch_decode(generated_ids, skip_special_tokens=False)[0]
        parsed = processor.post_process_generation(result, task="<OCR_WITH_REGION>", image_size=(image.width, image.height))
        return parsed.get("<OCR_WITH_REGION>", {}).get("text", str(parsed))

    # Generic VLM path
    _, processor, model, device = loaded
    # Build chat message with image
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "image", "image": image},
                {"type": "text", "text": OCR_PROMPT},
            ],
        }
    ]
    # Try chat template approach (Qwen2-VL style)
    try:
        text_prompt = processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = processor(
            text=[text_prompt],
            images=[image],
            return_tensors="pt",
            padding=True,
        ).to(device)
        output_ids = model.generate(**inputs, max_new_tokens=4096)
        # Trim input tokens
        generated_ids = [o[len(i):] for i, o in zip(inputs.input_ids, output_ids)]
        return processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    except Exception as e:
        logger.warning(f"Chat template failed ({e}), falling back to direct inputs")
        # Fallback: direct vision-language encoding
        inputs = processor(images=image, text=OCR_PROMPT, return_tensors="pt").to(device)
        output = model.generate(**inputs, max_new_tokens=4096)
        return processor.decode(output[0], skip_special_tokens=True)


# ─────────────────────────────────────────────────────────────────────────────
# Output formatters
# ─────────────────────────────────────────────────────────────────────────────

def _build_markdown(pages: list[dict]) -> str:
    parts = []
    for i, page in enumerate(pages, 1):
        if len(pages) > 1:
            parts.append(f"---\n\n## Page {i} — {page.get('source', '')}\n\n")
        parts.append(page["content"])
        parts.append("\n\n")
    return "".join(parts)


def _build_html(pages: list[dict]) -> str:
    import markdown as md_lib
    body_parts = []
    for i, page in enumerate(pages, 1):
        md_content = page["content"]
        html_content = md_lib.markdown(md_content, extensions=["tables", "fenced_code"])
        if len(pages) > 1:
            body_parts.append(f'<section class="page" id="page-{i}">')
            body_parts.append(f'<h2 class="page-header">Page {i} — {page.get("source", "")}</h2>')
            body_parts.append(html_content)
            body_parts.append("</section>")
        else:
            body_parts.append(html_content)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>OCR Result</title>
  <style>
    :root {{
      --bg: #ffffff; --fg: #1a1a2e; --accent: #4361ee;
      --border: #e0e0e0; --page-bg: #f8f9ff;
    }}
    body {{ font-family: 'Segoe UI', system-ui, sans-serif; max-width: 900px;
            margin: 0 auto; padding: 2rem; background: var(--bg); color: var(--fg); }}
    .page {{ background: var(--page-bg); border: 1px solid var(--border);
             border-radius: 8px; padding: 2rem; margin-bottom: 2rem; }}
    .page-header {{ color: var(--accent); border-bottom: 2px solid var(--accent);
                   padding-bottom: .5rem; }}
    table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
    th, td {{ border: 1px solid var(--border); padding: .5rem .75rem; text-align: left; }}
    th {{ background: var(--accent); color: white; }}
    tr:nth-child(even) td {{ background: #f0f4ff; }}
    blockquote {{ border-left: 4px solid var(--accent); margin: 0; padding-left: 1rem;
                  color: #555; font-style: italic; }}
    code {{ background: #f4f4f4; padding: .2em .4em; border-radius: 3px; font-size: .9em; }}
    pre code {{ display: block; padding: 1rem; overflow-x: auto; }}
  </style>
</head>
<body>
{''.join(body_parts)}
</body>
</html>"""


def _build_docx(pages: list[dict]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style the document
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, page in enumerate(pages, 1):
        if i > 1:
            doc.add_page_break()

        if len(pages) > 1:
            h = doc.add_heading(f"Page {i} — {page.get('source', '')}", level=2)
            h.runs[0].font.color.rgb = RGBColor(0x43, 0x61, 0xEE)

        content = page["content"]
        lines = content.split("\n")
        in_table = False
        table_rows: list[list[str]] = []

        def flush_table():
            nonlocal in_table, table_rows
            if not table_rows:
                return
            # Find max cols
            max_cols = max(len(r) for r in table_rows)
            t = doc.add_table(rows=len(table_rows), cols=max_cols)
            t.style = "Table Grid"
            for ri, row_data in enumerate(table_rows):
                for ci, cell_text in enumerate(row_data):
                    cell = t.rows[ri].cells[ci]
                    cell.text = cell_text.strip()
                    if ri == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            doc.add_paragraph("")
            in_table = False
            table_rows.clear()

        for line in lines:
            stripped = line.strip()

            # Table detection
            if re.match(r"\|.+\|", stripped):
                # Skip separator rows like |---|---|
                if re.match(r"[\|\-\s:]+$", stripped):
                    continue
                in_table = True
                cols = [c for c in stripped.split("|") if c.strip() != ""]
                table_rows.append(cols)
                continue
            else:
                if in_table:
                    flush_table()

            if not stripped:
                doc.add_paragraph("")
                continue

            # Headings
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            # Diagram / Image annotations
            elif re.match(r"\[(DIAGRAM|IMAGE|FIGURE):", stripped, re.IGNORECASE):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            # Bullet lists
            elif stripped.startswith(("- ", "* ", "• ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            else:
                # Inline bold/italic handling (basic)
                p = doc.add_paragraph()
                _add_formatted_run(p, stripped)

        if in_table:
            flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_run(paragraph: Any, text: str) -> None:
    """Parse basic markdown bold/italic and add formatted runs."""
    import re
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("__") and token.endswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


# ─────────────────────────────────────────────────────────────────────────────
# Core OCR pipeline
# ─────────────────────────────────────────────────────────────────────────────

async def _run_ocr_job(
    job_id: str,
    model_id: str,
    file_paths: list[Path],
    output_formats: list[str],
    db_session: Any = None,
) -> None:
    job = _jobs[job_id]
    job["status"] = "processing"
    pages: list[dict] = []

    try:
        # Gather all images with source labels
        all_images: list[tuple[str, Image.Image]] = []

        for fp in file_paths:
            suffix = fp.suffix.lower()
            if suffix == ".pdf":
                job["message"] = f"Converting PDF {fp.name}…"
                imgs = await asyncio.get_event_loop().run_in_executor(
                    None, _extract_images_from_pdf, fp
                )
                for i, img in enumerate(imgs, 1):
                    all_images.append((f"{fp.name} p.{i}", img))
            elif suffix == ".zip":
                job["message"] = f"Extracting ZIP {fp.name}…"
                pairs = await asyncio.get_event_loop().run_in_executor(
                    None, _extract_images_from_zip, fp
                )
                all_images.extend(pairs)
            else:
                img = await asyncio.get_event_loop().run_in_executor(None, _load_image, fp)
                all_images.append((fp.name, img))

        total = len(all_images)
        job["total_pages"] = total

        for idx, (source, img) in enumerate(all_images):
            job["message"] = f"Processing page {idx + 1}/{total}: {source}"
            job["progress"] = int((idx / total) * 85)

            raw = await asyncio.get_event_loop().run_in_executor(
                None, _run_inference, model_id, img
            )
            processed = _classify_elements(raw)
            pages.append({"source": source, "content": processed})
            job["pages"] = pages

        job["progress"] = 90
        job["message"] = "Generating output files…"

        # Build output files
        out_dir = OUTPUT_DIR / job_id
        out_dir.mkdir(parents=True, exist_ok=True)
        result_paths = []

        formats = [f.lower() for f in output_formats]

        md_content = _build_markdown(pages)

        if "md" in formats or "markdown" in formats:
            md_path = out_dir / "result.md"
            async with aiofiles.open(md_path, "w", encoding="utf-8") as f:
                await f.write(md_content)
            result_paths.append(str(md_path))

        if "html" in formats:
            html_content = await asyncio.get_event_loop().run_in_executor(
                None, _build_html, pages
            )
            html_path = out_dir / "result.html"
            async with aiofiles.open(html_path, "w", encoding="utf-8") as f:
                await f.write(html_content)
            result_paths.append(str(html_path))

        if "docx" in formats or "word" in formats:
            docx_bytes = await asyncio.get_event_loop().run_in_executor(
                None, _build_docx, pages
            )
            docx_path = out_dir / "result.docx"
            async with aiofiles.open(docx_path, "wb") as f:
                await f.write(docx_bytes)
            result_paths.append(str(docx_path))

        job["result_paths"] = result_paths
        job["status"] = "done"
        job["progress"] = 100
        job["message"] = "Done"
        job["completed_at"] = datetime.utcnow().isoformat()

        # Persist to DB
        if ENABLE_DB and db_session:
            from sqlalchemy import select
            result = await db_session.execute(
                select(JobRecord).where(JobRecord.job_id == job_id)
            )
            record = result.scalar_one_or_none()
            if record:
                record.status = "done"
                record.result_paths = json.dumps(result_paths)
                record.completed_at = datetime.utcnow()
                await db_session.commit()

    except Exception as exc:
        logger.exception(f"Job {job_id} failed: {exc}")
        job["status"] = "error"
        job["error"] = str(exc)
        job["message"] = f"Error: {exc}"
        job["progress"] = 0

        if ENABLE_DB and db_session:
            from sqlalchemy import select
            result = await db_session.execute(
                select(JobRecord).where(JobRecord.job_id == job_id)
            )
            record = result.scalar_one_or_none()
            if record:
                record.status = "error"
                record.error = str(exc)
                await db_session.commit()


# ─────────────────────────────────────────────────────────────────────────────
# App lifespan
# ─────────────────────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    logger.info("OCR App started. DB enabled: " + str(ENABLE_DB))
    yield
    logger.info("OCR App shutting down")


app = FastAPI(
    title="OCR Converter API",
    description="Multi-model OCR with layout understanding",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve output files
app.mount("/files", StaticFiles(directory=str(OUTPUT_DIR)), name="files")


# ─────────────────────────────────────────────────────────────────────────────
# API Routes
# ─────────────────────────────────────────────────────────────────────────────

# ── Models ──────────────────────────────────────────────────────────────────

@app.get("/api/models")
async def list_models():
    """List all available OCR models with download status."""
    return {"models": load_models()}


@app.post("/api/models/{model_id}/download")
async def start_download(model_id: str, background_tasks: BackgroundTasks):
    """Start downloading a model in the background."""
    model = get_model(model_id)
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")

    if _download_progress.get(model_id, {}).get("status") == "downloading":
        return {"status": "already_downloading"}

    background_tasks.add_task(download_model, model_id)
    _download_progress[model_id] = {
        "status": "queued",
        "progress": 0,
        "message": "Queued for download",
        "model_id": model_id,
    }
    return {"status": "started", "model_id": model_id}


@app.get("/api/models/{model_id}/download/progress")
async def download_progress_sse(model_id: str):
    """SSE stream for download progress."""
    async def event_generator():
        while True:
            prog = _download_progress.get(model_id, {
                "status": "idle", "progress": 0, "message": "Not started"
            })
            yield {"data": json.dumps(prog)}
            if prog.get("status") in ("done", "error"):
                break
            await asyncio.sleep(1)

    return EventSourceResponse(event_generator())


@app.get("/api/models/{model_id}/download/status")
async def download_status(model_id: str):
    """Poll download progress (non-SSE alternative)."""
    return _download_progress.get(model_id, {
        "status": "idle", "progress": 0, "message": "Not started", "model_id": model_id
    })


# ── OCR Jobs ────────────────────────────────────────────────────────────────

@app.post("/api/ocr/submit")
async def submit_ocr(
    background_tasks: BackgroundTasks,
    model_id: str = Form(...),
    output_formats: str = Form(...),   # comma-separated: "md,html,docx"
    files: list[UploadFile] = File(...),
    db: Any = Depends(get_db),
):
    """
    Submit OCR job.
    - model_id: one of the model IDs from /api/models
    - output_formats: comma-separated e.g. "md,html,docx"
    - files: images, PDFs, or ZIPs
    """
    model = get_model(model_id)
    if not model:
        raise HTTPException(status_code=404, detail=f"Model '{model_id}' not found")

    model_path = MODELS_DIR / model_id
    if not (model_path.exists() and any(model_path.iterdir())):
        raise HTTPException(
            status_code=409,
            detail=f"Model '{model_id}' is not downloaded. Download it first.",
        )

    formats = [f.strip().lower() for f in output_formats.split(",") if f.strip()]
    if not formats:
        raise HTTPException(status_code=400, detail="No output formats specified")

    valid_formats = {"md", "markdown", "html", "docx", "word"}
    invalid = set(formats) - valid_formats
    if invalid:
        raise HTTPException(status_code=400, detail=f"Invalid formats: {invalid}")

    # Validate and save uploaded files
    job_id = str(uuid.uuid4())
    job_upload_dir = UPLOAD_DIR / job_id
    job_upload_dir.mkdir(parents=True, exist_ok=True)
    saved_paths: list[Path] = []

    for upload in files:
        if not upload.filename:
            continue
        size = 0
        dest = job_upload_dir / upload.filename
        async with aiofiles.open(dest, "wb") as out_file:
            while chunk := await upload.read(1024 * 1024):
                size += len(chunk)
                if size > MAX_UPLOAD_MB * 1024 * 1024:
                    raise HTTPException(
                        status_code=413,
                        detail=f"File {upload.filename} exceeds {MAX_UPLOAD_MB}MB limit",
                    )
                await out_file.write(chunk)
        # Validate type
        suffix = dest.suffix.lower()
        allowed = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif", ".pdf", ".zip"}
        if suffix not in allowed:
            dest.unlink()
            raise HTTPException(status_code=415, detail=f"Unsupported file type: {suffix}")
        saved_paths.append(dest)

    if not saved_paths:
        raise HTTPException(status_code=400, detail="No valid files uploaded")

    # Create job
    job = _new_job(job_id, model_id, formats)
    job["input_files"] = [p.name for p in saved_paths]

    # Persist to DB
    if ENABLE_DB and db is not None:
        from sqlalchemy import select
        record = JobRecord(
            job_id=job_id,
            status="pending",
            model_id=model_id,
            input_files=json.dumps([p.name for p in saved_paths]),
            output_formats=json.dumps(formats),
        )
        db.add(record)
        await db.commit()

    # Start background processing
    background_tasks.add_task(_run_ocr_job, job_id, model_id, saved_paths, formats, db)

    return {"job_id": job_id, "status": "pending"}


@app.get("/api/ocr/jobs/{job_id}")
async def get_job_status(job_id: str):
    """Get job status and result paths."""
    job = _jobs.get(job_id)
    if not job:
        # Try DB
        if ENABLE_DB:
            raise HTTPException(status_code=404, detail="Job not found (not in memory; check DB)")
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@app.get("/api/ocr/jobs/{job_id}/progress")
async def job_progress_sse(job_id: str):
    """SSE stream for OCR job progress."""
    async def event_generator():
        while True:
            job = _jobs.get(job_id)
            if not job:
                yield {"data": json.dumps({"status": "not_found"})}
                break
            yield {"data": json.dumps({
                "status": job["status"],
                "progress": job["progress"],
                "message": job["message"],
                "total_pages": job.get("total_pages", 0),
            })}
            if job["status"] in ("done", "error"):
                break
            await asyncio.sleep(0.5)

    return EventSourceResponse(event_generator())


@app.get("/api/ocr/jobs/{job_id}/download/{filename}")
async def download_result(job_id: str, filename: str):
    """Download a specific result file."""
    job = _jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["status"] != "done":
        raise HTTPException(status_code=409, detail="Job not complete")

    file_path = OUTPUT_DIR / job_id / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    media_type, _ = mimetypes.guess_type(str(file_path))
    return FileResponse(
        str(file_path),
        media_type=media_type or "application/octet-stream",
        filename=filename,
    )


@app.get("/api/ocr/jobs")
async def list_jobs(limit: int = 50):
    """List recent jobs (in-memory)."""
    jobs = sorted(_jobs.values(), key=lambda j: j.get("created_at", ""), reverse=True)
    return {"jobs": jobs[:limit]}


@app.delete("/api/ocr/jobs/{job_id}")
async def delete_job(job_id: str):
    """Delete a job and its files."""
    if job_id not in _jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    _jobs.pop(job_id, None)
    for d in (UPLOAD_DIR / job_id, OUTPUT_DIR / job_id):
        if d.exists():
            shutil.rmtree(d)
    return {"deleted": job_id}


# ── History (DB only) ────────────────────────────────────────────────────────

@app.get("/api/history")
async def get_history(limit: int = 50, db: Any = Depends(get_db)):
    """Get persisted job history (requires ENABLE_DB=true)."""
    if not ENABLE_DB or db is None:
        raise HTTPException(
            status_code=501,
            detail="History is disabled. Set ENABLE_DB=true in .env to enable.",
        )
    from sqlalchemy import select, desc
    result = await db.execute(
        select(JobRecord).order_by(desc(JobRecord.created_at)).limit(limit)
    )
    records = result.scalars().all()
    return {
        "history": [
            {
                "job_id": r.job_id,
                "status": r.status,
                "model_id": r.model_id,
                "input_files": json.loads(r.input_files or "[]"),
                "output_formats": json.loads(r.output_formats or "[]"),
                "created_at": r.created_at.isoformat() if r.created_at else None,
                "completed_at": r.completed_at.isoformat() if r.completed_at else None,
                "result_paths": json.loads(r.result_paths or "[]"),
                "error": r.error,
            }
            for r in records
        ]
    }


# ── Health ───────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "db_enabled": ENABLE_DB,
        "models_dir": str(MODELS_DIR),
        "timestamp": datetime.utcnow().isoformat(),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)