"""
formatters.py — Convert OCRResult into the requested output formats.

Supported formats: md | html | docx
"""

from __future__ import annotations

import os
import re
import textwrap
from pathlib import Path

from ocr_engine import OCRResult, RegionType


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def to_markdown(result: OCRResult, title: str = "OCR Output") -> str:
    lines = [f"# {title}\n"]
    for p in result.pages:
        if len(result.pages) > 1:
            lines.append(f"\n---\n## Page {p.page_number}\n")
        for r in p.regions:
            if r.region_type == RegionType.HEADER:
                lines.append(f"## {r.content}\n")
            elif r.region_type == RegionType.TABLE:
                lines.append(r.content + "\n")
            elif r.region_type == RegionType.FORMULA:
                lines.append(f"$$\n{r.content}\n$$\n")
            elif r.region_type == RegionType.DIAGRAM:
                lines.append(f"> 📊 **[Diagram / Figure]**\n>\n> {r.content}\n")
            elif r.region_type == RegionType.IMAGE:
                lines.append(f"> 🖼️ **[Embedded Image]**\n>\n> {r.content}\n")
            elif r.region_type == RegionType.CAPTION:
                lines.append(f"*{r.content}*\n")
            elif r.region_type == RegionType.FOOTER:
                lines.append(f"---\n_{r.content}_\n")
            else:
                lines.append(r.content + "\n")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def to_html(result: OCRResult, title: str = "OCR Output") -> str:
    body_parts: list[str] = []

    for p in result.pages:
        if len(result.pages) > 1:
            body_parts.append(f'<section class="page" id="page-{p.page_number}">')
            body_parts.append(f'<h2 class="page-label">Page {p.page_number}</h2>')
        for r in p.regions:
            if r.region_type == RegionType.HEADER:
                body_parts.append(f'<h2 class="ocr-header">{_esc(r.content)}</h2>')
            elif r.region_type == RegionType.TABLE:
                body_parts.append(_md_table_to_html(r.content))
            elif r.region_type == RegionType.FORMULA:
                body_parts.append(
                    f'<div class="ocr-formula"><code>{_esc(r.content)}</code></div>'
                )
            elif r.region_type == RegionType.DIAGRAM:
                body_parts.append(
                    f'<figure class="ocr-diagram">'
                    f'<figcaption>📊 Diagram / Figure</figcaption>'
                    f'<p>{_esc(r.content)}</p>'
                    f'</figure>'
                )
            elif r.region_type == RegionType.IMAGE:
                body_parts.append(
                    f'<figure class="ocr-image">'
                    f'<figcaption>🖼 Embedded Image</figcaption>'
                    f'<p>{_esc(r.content)}</p>'
                    f'</figure>'
                )
            elif r.region_type == RegionType.CAPTION:
                body_parts.append(f'<p class="ocr-caption"><em>{_esc(r.content)}</em></p>')
            elif r.region_type == RegionType.FOOTER:
                body_parts.append(f'<footer class="ocr-footer"><small>{_esc(r.content)}</small></footer>')
            else:
                paragraphs = r.content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        body_parts.append(f'<p class="ocr-text">{_esc(para.strip())}</p>')
        if len(result.pages) > 1:
            body_parts.append("</section>")

    body = "\n".join(body_parts)

    return textwrap.dedent(f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{_esc(title)}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; }}
    body {{
      font-family: 'Segoe UI', system-ui, sans-serif;
      max-width: 860px;
      margin: 2rem auto;
      padding: 0 1.5rem;
      color: #1a1a1a;
      line-height: 1.7;
    }}
    h1 {{ font-size: 2rem; border-bottom: 3px solid #4f46e5; padding-bottom: .5rem; }}
    h2 {{ font-size: 1.4rem; color: #4f46e5; }}
    .page {{ border-left: 4px solid #e0e7ff; padding-left: 1.2rem; margin: 2rem 0; }}
    .page-label {{ color: #6366f1; font-size: 1rem; text-transform: uppercase; letter-spacing: .08em; }}
    .ocr-table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
    .ocr-table th, .ocr-table td {{ border: 1px solid #d1d5db; padding: .5rem .75rem; }}
    .ocr-table th {{ background: #f3f4f6; font-weight: 600; }}
    .ocr-formula {{ background: #fefce8; border: 1px solid #fde047; border-radius: 4px;
                    padding: .75rem 1rem; font-family: monospace; overflow-x: auto; }}
    .ocr-diagram, .ocr-image {{ background: #f0fdf4; border: 1px solid #86efac;
                                 border-radius: 6px; padding: 1rem; margin: 1rem 0; }}
    .ocr-diagram figcaption, .ocr-image figcaption {{ font-weight: 700; margin-bottom: .4rem; }}
    .ocr-caption {{ color: #6b7280; font-size: .9rem; }}
    .ocr-footer {{ border-top: 1px solid #e5e7eb; margin-top: 2rem; color: #9ca3af; }}
    .ocr-text {{ margin: .5rem 0; }}
  </style>
</head>
<body>
  <h1>{_esc(title)}</h1>
  {body}
</body>
</html>
""")


def _esc(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )


def _md_table_to_html(md_table: str) -> str:
    """Convert a markdown-style table to HTML table."""
    lines = [l.strip() for l in md_table.strip().splitlines() if l.strip()]
    if not lines:
        return f'<pre class="ocr-table-raw">{_esc(md_table)}</pre>'

    rows = []
    header_done = False
    html = ['<table class="ocr-table">']
    for line in lines:
        if re.match(r"^[\|\s\-:]+$", line):
            header_done = True
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        tag = "th" if not header_done else "td"
        html.append("<tr>" + "".join(f"<{tag}>{_esc(c)}</{tag}>" for c in cells) + "</tr>")
        if not header_done:
            header_done = True  # first data row is header if no separator yet
    html.append("</table>")
    return "\n".join(html)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def to_docx(result: OCRResult, output_path: Path, title: str = "OCR Output") -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in result.pages:
        if len(result.pages) > 1:
            doc.add_page_break()
            heading = doc.add_heading(f"Page {p.page_number}", level=1)

        for r in p.regions:
            if r.region_type == RegionType.HEADER:
                doc.add_heading(r.content, level=2)

            elif r.region_type == RegionType.TABLE:
                _add_md_table_to_docx(doc, r.content)

            elif r.region_type == RegionType.FORMULA:
                para = doc.add_paragraph()
                run = para.add_run(r.content)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                para.style = doc.styles["No Spacing"]
                # Light yellow shading via XML
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "FEFCE8")
                pPr.append(shd)

            elif r.region_type in (RegionType.DIAGRAM, RegionType.IMAGE):
                label = "📊 [Diagram / Figure]" if r.region_type == RegionType.DIAGRAM else "🖼 [Embedded Image]"
                para = doc.add_paragraph()
                run = para.add_run(label)
                run.bold = True
                run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
                if r.content.strip():
                    doc.add_paragraph(r.content)

            elif r.region_type == RegionType.CAPTION:
                para = doc.add_paragraph(r.content)
                para.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]

            elif r.region_type == RegionType.FOOTER:
                para = doc.add_paragraph(r.content)
                run = para.runs[0] if para.runs else para.add_run(r.content)
                run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
                run.font.size = Pt(9)

            else:
                paragraphs = r.content.split("\n")
                for text in paragraphs:
                    if text.strip():
                        doc.add_paragraph(text.strip())

    doc.save(str(output_path))


def _add_md_table_to_docx(doc, md_table: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    lines = [l.strip() for l in md_table.strip().splitlines() if l.strip()]
    data_rows = [
        [c.strip() for c in l.strip("|").split("|")]
        for l in lines
        if not re.match(r"^[\|\s\-:]+$", l)
    ]

    if not data_rows:
        doc.add_paragraph(md_table)
        return

    col_count = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(data_rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < col_count:
                cell = row.cells[col_idx]
                cell.text = cell_text
                if row_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

SUPPORTED_FORMATS = {"md", "html", "docx"}


def export(
    result: OCRResult,
    formats: list[str],
    output_dir: Path,
    base_name: str,
    title: str = "OCR Output",
) -> dict[str, Path]:
    """
    Export OCRResult to all requested formats.

    Returns a dict mapping format → output file path.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    outputs: dict[str, Path] = {}

    for fmt in formats:
        fmt = fmt.lower().strip()
        if fmt not in SUPPORTED_FORMATS:
            continue

        if fmt == "md":
            path = output_dir / f"{base_name}.md"
            path.write_text(to_markdown(result, title=title), encoding="utf-8")
            outputs["md"] = path

        elif fmt == "html":
            path = output_dir / f"{base_name}.html"
            path.write_text(to_html(result, title=title), encoding="utf-8")
            outputs["html"] = path

        elif fmt == "docx":
            path = output_dir / f"{base_name}.docx"
            to_docx(result, path, title=title)
            outputs["docx"] = path

    return outputs
